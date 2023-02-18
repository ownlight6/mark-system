# -*- coding: utf-8 -*-

import glob
# from models import Job
import json
import sys

import docx
from django.http import JsonResponse
from django.shortcuts import render
from docx import Document
from PIL import Image
from io import BytesIO
import base64

from mydjango5 import settings

reload(sys)
sys.setdefaultencoding("utf-8")
import os
import re


# Create your views here.

# 登录后的首页
def indexPage(req):
    username = req.COOKIES.get('cookie_username', '')
    response = render(req, 'index.html', {'username': username})
    response['Cache-Control'] = 'no-cache'
    response['X-Frame-Options'] = 'ALLOW-FROM'

    return response

    # return render(req,'index.html')


# 登录页面
def loginPage(request):
    response = render(request, 'loginPage.html')
    response['Cache-Control'] = 'no-cache'
    response['X-Frame-Options'] = 'ALLOW-FROM'
    # response.setHeader("x-frame-options", "ALLOW-FROM")
    return response


# 获取目录初始化树
path = "D:\mytree"
parent = 0
Id = 0

# 当前环境  本地2 服务器1
now_server = 2

# @login_required
def get_txtcatalog(request):
    # path = "D:/mytree"
    path = settings.UPLOAD_URL
    jsonstr = "["
    jsonstr = funForGetFileCatalog(jsonstr, path, 0)
    jsonstr += "]"
    # 部署服务器
    if now_server == 1:
        jsonstr = jsonstr.encode("utf-8")
    # 本地
    else:
        jsonstr = jsonstr.decode("gbk").encode("utf-8")
    resaultCatalog = eval(jsonstr)

    return JsonResponse({'catalog': json.dumps(resaultCatalog)})


# 递归读取本地文件目录函数
def funForGetFileCatalog(jsonstr, path, parent):
    global Id
    global count
    # root:false 用来判断是否是顶层文件夹，只能判断两层，往下第三层会出错
    for i, fn in enumerate(glob.glob(path + os.sep + '*')):

        if os.path.isdir(fn):
            jsonstr += '{"attributes":{"id":"' + str(Id) + '"},"parent":"' + str(
                parent) + '","root":"true","name":"' + os.path.basename(fn) + '","children":['
            parent = Id
            Id += 1
            for j, li in enumerate(glob.glob(fn + os.sep + '*')):

                if os.path.isdir(li):
                    jsonstr += '{"attributes":{"id":"' + str(Id) + '"},"parent":"' + str(
                        parent) + '","root":"false","name":"' + os.path.basename(
                        li) + '","children":['
                    parent = Id
                    Id += 1
                    funForGetFileCatalog(li, parent)
                    jsonstr += "]}"
                    if j < len(glob.glob(fn + os.sep + '*')) - 1:
                        jsonstr += ","

                else:
                    # if (li.split('.')[-1] == 'docx'):
                    # 暂时注释掉
                    jsonstr += '{"attributes":{"id":"' + str(Id) + '"},"parent":"' + str(
                        parent) + '","root":"false","name":"''' + os.path.basename(
                        li) + '","type":"leaf"}'
                    Id += 1
                    if j < len(glob.glob(fn + os.sep + '*')) - 1:
                        jsonstr += ","
            jsonstr += "]}"
            if i < len(glob.glob(path + os.sep + '*')) - 1:
                jsonstr += ","

        else:
            if (fn.split('.')[-1] == 'txt'):
                jsonstr += '{"attributes":{"id":"' + str(Id) + '"},"parent":"' + str(
                    parent) + '","root":"false","name":"''' + os.path.basename(fn) + '","type":"leaf"}'
                Id += 1
                if i < len(glob.glob(path + os.sep + '*')) - 1:
                    jsonstr += ","
    return jsonstr


# 得到所点击txt的文本内容
# @login_required
def get_text(request):
    path = settings.UPLOAD_URL
    pathName = request.POST["pathname"]
    lastAnnName = request.POST["lastAnnName"]

    # if(lastAnnName!=""):
    #     os.remove(lastAnnName)

    #服务器
    if now_server == 1:
        completeFilePath = path + '/' + pathName
    else:
        completeFilePath = path + '/' + pathName.decode('utf-8').encode("gbk")
    pos = completeFilePath.rfind("/")

    completeConfPath = completeFilePath[:pos] + '/annotation.conf'
    #服务器
    if now_server == 1:
        get_pictures(completeFilePath,  "temp/"+pathName.split("/")[1])
    else:
        get_pictures(completeFilePath,  "temp/"+pathName.decode('utf-8').encode("gbk").split("/")[1])

    if completeFilePath.split(".")[-1].encode("utf-8") == "txt":
        # 读取txt文件内容
        f = open(completeFilePath)
        stringaa = ""
        textResault = f.read()
        # print textResault
        # for name in textResault:
        #     name = name.decode('gbk').encode('utf-8')
        # stringaa = stringaa + name
        # textResault = textResault.decode('gbk').encode('utf-8')
        # print textResault
        f.close()

        stringTextResault = ""
        textResault = textResault.replace("\n", "\\n")

        stringTextResault = "{'text':" + "\"" + textResault + "\"" + "}"
        # print("ssss",stringTextResault)
        jsonTextResault = eval(stringTextResault)
        # # 读取配置文件内容
        # stringConfResault = getConfComment(completeConfPath)
        # dictConfResault = eval(stringConfResault)

        return JsonResponse(jsonTextResault)

    if completeFilePath.split(".")[-1].encode("utf-8") == "docx":

        img_arr = []
        file_path = completeFilePath.split("/")[2]
        if os.path.exists("temp/"+file_path):
            for file_name in os.listdir("temp/"+file_path):
                if file_name is not None:
                    img_arr.append("temp/"+file_path+"/"+file_name)

        # 读取doc文档内容
        docxString = []
        imgIndex = 0
        document = Document(completeFilePath)
        for paragraph in document.paragraphs:
            docxString.append(paragraph.text)
            image_list = paragraph._element.xpath('.//pic:pic')
            if not image_list:
                continue
            for image in image_list:
                if image is not None:
                    imgstring = base64.b64encode(open(img_arr[imgIndex], 'rb').read())
                    docxString.append(imgstring)
                    imgIndex = imgIndex+1

        # print docxString
        # stringTextResault = "{'text':" + "\"" + docxString + "\"" + "}"
        # print stringTextResault
        dict_text = {'text': docxString}
        # jsonTextResault = eval(stringTextResault)
        return JsonResponse(dict_text)
    return ""


def get_pictures(word_path, result_path):
    """
    图片提取
    :param word_path: word路径
    :param result_path: 结果路径
    :return:
    """
    doc = docx.Document(word_path)
    dict_rel = doc.part._rels
    for rel in dict_rel:
        rel = dict_rel[rel]
        if "image" in rel.target_ref:
            if not os.path.exists(result_path):
                os.makedirs(result_path)
            img_name = re.findall("/(.*)", rel.target_ref)[0]
            word_name = os.path.splitext(word_path)[0]
            if os.sep in word_name:
                new_name = word_name.split('\\')[-1]
            else:
                new_name = word_name.split('/')[-1]
            img_name = new_name + "_" + img_name
            #服务器
            if now_server == 1:
                img_name = img_name.split("/")[2]
            with open(result_path + "/" + img_name, "wb") as f:
                f.write(rel.target_part.blob)


# 读取所选txt的已有标注文件目录
# @login_required
def get_annCatalog(request):
    path = settings.UPLOAD_URL
    pathName = request.POST["pathname"]
    # print pathName
    completeFilePath = path + '/' + pathName

    pos = completeFilePath.rfind("/")

    # txtName = completeFilePath[pos+1:-4]
    docxName = completeFilePath[pos + 1:-5]
    annPath = completeFilePath[:pos]
    # print docxName
    annjsonstr = funForGetAnnCatalog(annPath, docxName, '')
    annjsonstr += "]"
    annjsonstr = annjsonstr.replace('},]', '}]')
    # print(annjsonstr)
    resaultAnnCatalog = eval(annjsonstr)
    return JsonResponse({'annCatalog': json.dumps(resaultAnnCatalog, ensure_ascii=False)})


# 读取所选txt的已有标注文本函数
parent = 0
annid = 1
s = "不打开已有标注"
s_unicode = s.decode("utf-8")
annjsonstr = ''


def funForGetAnnCatalog(path, filename, username):
    global annid
    global annjsonstr
    annjsonstr = '[{"attributes":{"id":"0"},"name":' + '"' + s_unicode + '"' + '},'
    fileTheName = filename
    # print fileTheName
    # root:false 用来判断是否是顶层文件夹，只能判断两层，往下第三层会出错
    for i, fn in enumerate(glob.glob(path + os.sep + '*.ann')):
        # print fn
        if (fileTheName in fn):
            annjsonstr += '{"attributes":{"id":"' + str(annid) + '"},"name":"''' + os.path.basename(fn) + '"}'
            annid += 1
            if i < len(glob.glob(path + os.sep + '*.ann')) - 1:
                annjsonstr += ","
    return annjsonstr


# 读取配置文件
# @login_required
def get_conf(request):
    path = settings.UPLOAD_URL
    pathName = request.POST["pathname"]
    completeFilePath = path + '/' + pathName

    pos = completeFilePath.rfind("/")

    completeConfPath = completeFilePath[:pos] + '/annotation.conf'

    # 读取配置文件内容构造json格式
    stringConfResault = getConfComment(completeConfPath)
    dictConfResault = eval(stringConfResault)
    # print dictConfResault
    return JsonResponse({'conf': json.dumps(dictConfResault)})


# 分割conf文本
def parse_conf(str):
    return str.split(',')


# 读取配置文件构造json数据
def getConfComment(confPath):
    entity = '[entities]:'
    events = '[events]:'
    confJsonstr = '['
    with open(confPath, "r") as f:  # 该路径需要从前端获得
        lines = f.readlines()
        for line in lines:
            line = line[:-1]
            if entity in line:
                confJsonstr += '{"entity":['
                value = parse_conf(line.split(':')[1])
                for temp in value:
                    confJsonstr += '{"text":"' + temp + '"},'
                confJsonstr += ']'
            if events in line:
                confJsonstr += ',"events":['
                value = parse_conf(line.split(':')[1])
                for temp in value:
                    confJsonstr += '{"text":"' + temp + '"},'
                confJsonstr += ']'

    confJsonstr += '}]'
    confJsonstr = confJsonstr.replace(',]', ']')
    return confJsonstr


def uploadPage(request):
    return render(request, 'upload.html')


def upload(request):
    if request.method == "POST":
        # print(request.FILES)
        # upload_Form=request.
        # formdata = request.POST
        # handle_upload_file(request.POST['fafafa'], str(request.POST['fafafa']))
        # handle_upload_file(request.FILES['fafafa'], str(request.FILES['fafafa']))
        handle_upload_file(request.FILES['file'], str(request.FILES['file']))
        return JsonResponse({'msg': '上传成功！'})  # 此处简单返回一个成功的消息，在实际应用中可以返回到指定的页面中

    return JsonResponse({'msg': '上传失败！'})


def handle_upload_file(file, filename):
    path = 'userUploadFile/file and filefolder/'  # 上传文件的保存路径，可以自己指定任意的路径
    filename = filename.decode('utf-8')
    if not os.path.exists(path):
        os.makedirs(path)
    with open(path + filename, 'wb+') as destination:
        for chunk in file.chunks():
            destination.write(chunk)


def save_QA(request):
    path = settings.UPLOAD_URL

    question = request.POST["question"]
    answer = request.POST["answer"]

    # if(lastAnnName!=""):
    #     os.remove(lastAnnName)

    document = Document(path + '/' + 'file and filefolder/QA.docx')
    for paragraph in document.paragraphs:
        paragraph.clear()
    p1 = document.add_paragraph(question)
    p2 = document.add_paragraph(answer)
    document.save(path + '/' + 'file and filefolder/QA.docx')
    stringTextResault = "{'code':" + "1" + ",'msg':'保存QA成功！'" + "}"
    jsonTextResault = eval(stringTextResault)
    return JsonResponse(jsonTextResault)
